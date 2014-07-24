from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_UNDERLINE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
import shutil

class DocBuilder(object):
	"""docstring for DocBuilder"""
	def __init__(self):
		self.document = Document()
		self.EXPORT_DIRECTORY = 'StatementExports'

		# Create export directory if it doesn't exist
		if not os.path.exists(self.EXPORT_DIRECTORY):
			os.makedirs(self.EXPORT_DIRECTORY)

	def createTable(self, numberOfColumns, tableData):
		"""Expects tableData to be a list of list of Cell objects"""
		self.table = self.document.add_table(rows=0, cols=0)

		# Contains the column objects whose widths can be set later on
		self.columns = []

		# Write the table headers
		for x in xrange(numberOfColumns):
			self.columns.append(self.table.add_column())

		# Write the table data
		for data in tableData:
			rowCells = self.table.add_row().cells
			for i in xrange(0, len(data)):
				cellData = data[i]
				cell = rowCells[i]
				cell.text = cellData.text
				if 'underline' in cellData.tags:
					cell.paragraphs[0].runs[0].underline = True
				if 'double_underline' in cellData.tags:
					cell.paragraphs[0].runs[0].underline = WD_UNDERLINE.DOUBLE
				if 'bold' in cellData.tags:
					cell.paragraphs[0].runs[0].bold = True

		# Set column widths to be even by default; can be set otherwise afterwards
		width = Inches(6/len(self.columns))
		for column in self.columns:
			column.width = width

	def createHeading(self, text, align = "center"):
		"""Possible values for align: "center" "left" "right" """
		self.heading = self.document.add_heading(text)

		if align == "center":
			alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
		elif align == "left":
			alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
		elif align == "right":
			alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
			
		self.heading.alignment = alignment

	def save(self, fileName):
		"""Saves the document"""
		self.fileName = fileName
		self.document.save(fileName)
		shutil.move(self.fileName, self.EXPORT_DIRECTORY+'/'+self.fileName)


class CellData(object):
	"""Represents a cell in a docx table, where self.text is the cell text, 
	and self.tags is a list of strings to specify the format of that text e.g. "underlined" """
	def __init__(self, text, *args):
		self.tags = []
		for i, value in enumerate(args):
			self.tags.append(value)
		self.text = text


def resource_path(relative_path):
    """ Get absolute path to resource, works for dev and for PyInstaller """
    try:
        # PyInstaller creates a temp folder and stores path in _MEIPASS
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")

    return os.path.join(base_path, relative_path)
# for testing with pyinstaller
def createDummyDocument():
	document = Document(resource_path(os.path.join("assets","default.docx")))

	document.add_heading('Document Title', 0)

	p = document.add_paragraph('A plain paragraph having some ')
	p.add_run('bold').bold = True
	p.add_run(' and some ')
	p.add_run('italic.').italic = True

	document.add_heading('Heading, level 1', level=1)
	document.add_paragraph('Intense quote', style='IntenseQuote')

	document.add_paragraph(
	    'first item in unordered list', style='ListBullet'
	)
	document.add_paragraph(
	    'first item in ordered list', style='ListNumber'
	)

	# document.add_picture('monty-truth.png', width=Inches(1.25))

	table = document.add_table(rows=1, cols=0, style='LightShading-Accent1')
	column0 = table.add_column()
	column0.width = 5000000
	column1 = table.add_column()
	column2 = table.add_column()
	# print table.columns.cells
	table.columns.width = 200000

	hdr_cells = table.rows[0].cells
	hdr_cells[0].text = 'Qty'
	hdr_cells[1].text = 'Id'
	hdr_cells[2].text = 'Desc'

	tableData = [["5", "Apples", "yummy"], ["6", "Oranges", "they're orange"]]

	for row in tableData:
		row_cells = table.add_row().cells
		row_cells[0].text = row[0]
		row_cells[1].text = row[1]
		row_cells[2].text = row[2]


	document.add_page_break()

	document.save('demo.docx')

if __name__ == "__main__":
	docBuilder = DocBuilder()

	rowData = []
	rowData.append([CellData("Orange"), CellData("5"), CellData("Yummy", "underline")])
	rowData.append([CellData("Apple"), CellData("2"), CellData("They're red.", "underline", "bold")])

	docBuilder.createTable(3, rowData)
	docBuilder.columns[0].width = Inches(1)
	docBuilder.columns[1].width = Inches(1)
	docBuilder.columns[2].width = Inches(4)

	docBuilder.save("docbuildertest.docx")


