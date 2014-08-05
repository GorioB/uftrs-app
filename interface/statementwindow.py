from Tkinter import *
from ttk import *
from tkFont import Font
from textable import TextTable
from lib.floattostr import *
from ScrolledFrame import VerticalScrolledFrame
from lib.timeFuncs import *
from lib.app import *
from interface.docbuilder import DocBuilder, CellData
from docx.shared import Inches
from cashflowswindow import filterDOT
from lib.getStartingBalance import *
from lib.getOAL import *
from enumerateNotes import *

MONTHS=["","January","February","March","April","May","June","July","August","September","October","November","December"]
def tab(n=1):
	return "    "*n

def floatToStrParenNeg(n):
	if n<0:
		return "("+floatToStr(n).strip("-")+")"
	else:
		return floatToStr(n)
class StatementWindow(Frame,object):
	def __init__(self,parent,app,deletedVar,**kwargs):
		Frame.__init__(self,parent,**kwargs)
		self.app = app
		self.parent = parent
		self.initUI()

	def initUI(self):
		self.printFrame = Frame(self)
		self.printFrame.pack(fill=X,expand=0,side=TOP)
		self.printButton = Button(self.printFrame,text="Export to docx",command=self.exportCallback)
		self.printButton.pack(fill=NONE,expand=0,side=LEFT)
		self.mainFrame = VerticalScrolledFrame(self)
		self.mainFrame.pack(fill=BOTH,expand=1,side=TOP)
		self.headerField = Text(self.mainFrame.interior,bd=0,width=0,state='disabled',height=3)
		self.headerField.tag_configure("center",justify="center")
		self.boldFont = f = Font(self.headerField,self.headerField.cget('font'))
		f.configure(weight='bold')
		self.headerField.configure(font=f)
		self.headerField.pack(fill=X,expand=0,side=TOP)

		self.cashFlowsText = TextTable(self.mainFrame.interior,
			aligns=['left','left','right','right','right','right'],
			weights=[4,4,1,1,1,1],
			width=500)
		self.cashFlowsText.pack(expand=1,fill=BOTH,side=TOP)

		self.oalText = Text(self.mainFrame.interior,bd=0,width=0,state='disabled',height=0)
		self.oalText.tag_configure("bold",font=f)
		self.oalText.pack(fill=X,expand=0,side=TOP)

		self.noteChunk=NoteChunk(self.mainFrame.interior,self.app)
		self.noteChunk.pack(fill=BOTH,expand=1,side=TOP)
		self.noteBlocks = self.noteChunk.noteBlocks

	def populateTree(self):
		tStart = secsToDay(self.app.timeFrame[0]).split("-")
		sYear = tStart[0]
		sDay = tStart[2]
		sMonth = MONTHS[int(tStart[1])]
		tStart = str(sMonth)+" "+str(sDay)+", "+str(sYear)
		tEnd = secsToDay(self.app.timeFrame[-1]).split("-")
		month = MONTHS[int(tEnd[1])]
		year = tEnd[0]
		day = tEnd[2]
		self.lines=[]
		self.headerField['state']='normal'
		self.headerField.delete('1.0','end')
		self.headerField.insert('1.0',self.app.councilName+"\nStatement of Cash Flows\nFor the Semester Ended "+month+" "+day+", "+year)
		self.headerField.tag_add('center','1.0','end')
		#self.headerField['state']='disabled'

		self.lines.append([['CASH INFLOWS','Note','','','',''],[],[1,0,0,0,0,0]])
		cashFlowList = self.app.listCashflows(showDeleted=False)
		inflowList = [i for i in cashFlowList if i.source.content.split(":")[0]=="CashReceipt"]
		totalInflows=0
		self.firstInflow=0
		self.firstInflowTotal=0
		self.lineTypes="Inflow"
		#replace with db stuff
		inflowTypeList = ['Council Mandated Funds','General Sponsorship Inflows','Income Generating Projects','Other Inflows']
		for inflowType in inflowTypeList:
			partialList = filterDOT(self.app,[i for i in inflowList if i.getContents().category.content==inflowType])
			lines,partialTotal = self.getInflows(partialList)
			if lines!=[]:
				self.lines.append([[inflowType,'','','','',''],[],[]])
			self.lines=self.lines+lines
			totalInflows+=partialTotal

		self.firstInflow=0
		self.firstInflowTotal=0
		self.lineTypes="Outflow"
		self.lines.append([[tab(3)+"Total Inflows",'','','','P',floatToStr(totalInflows)],[0,0,0,0,0,1],[]])
		self.lines.append([["",'','','','',''],[0,0,0,0,0,0],[0,0,0,0,0,0]])
		self.lines.append([['CASH OUTFLOWS','','','','',''],[],[1,0,0,0,0,0]])
		outflowList = [i for i in cashFlowList if i.source.content.split(":")[0] in ("OME","COCPNote","LTINote","OONote")]
		totalOutflows=0
		self.firstOutflow=0
		self.firstOutflowTotal=0
		outflowNames ={"COCPNote":"Council and Other College Projects","LTINote":"Long Term Investment","OME":"Operation and Maintenance Expenses","OONote":"Other Outflows"}
		for outflowType in ("COCPNote","LTINote","OME","OONote"):
			partialList = filterDOT(self.app,[i for i in outflowList if i.source.content.split(":")[0]==outflowType])
			lines,partialTotal = self.getInflows(partialList)
			if lines!=[]:
				self.lines.append([[outflowNames[outflowType],'','','','',''],[],[]])
			self.lines=self.lines+lines
			totalOutflows+=partialTotal
		self.lines.append([[tab(3)+"Total Outflows",'','','','P',"("+floatToStr(totalOutflows)+")"],[0,0,0,0,0,1],[]])
		self.lines = self.setUnderlines(self.lines)
		totalFlows = totalInflows-totalOutflows
		strTotalFlows = floatToStrParenNeg(totalFlows)


		self.lines.append([["Net Cash Flow",'','','','',strTotalFlows],[0,0,0,0,0,0],[1,0,0,0,0,0]])
		#beginning balance filler
		beginningBalance=getStartBalance(self.app,int(self.app.timeFrame[0]))
		self.lines.append([["Add: Beginning Cash Balance, "+tStart,'','','','',floatToStr(beginningBalance)],[0,0,0,0,0,1],[0,0,0,0,0,0]])

		self.lines.append([["Ending Cash Balance, "+month+" "+day+", "+year,'','','','',floatToStrParenNeg(beginningBalance+totalFlows)],[0,0,0,0,0,1],[1,0,0,0,0,0]])

		self.lines.append([["",'','','','',''],[0,0,0,0,0,0],[0,0,0,0,0,0]])
		self.lines.append([["Cash Breakdown",'','','','',''],[0,0,0,0,0,0],[1,0,0,0,0,0]])
		self.lines.append([['Cash on hand','','','','',self.app.cashOnHand],[0,0,0,0,0,0],[0,0,0,0,0,0]])
		self.lines.append([['Cash in bank','','','','',self.app.cashInBank],[0,0,0,0,0,0],[0,0,0,0,0,0]])
		totalBalance = floatToStr(strToFloat(self.app.cashOnHand)+strToFloat(self.app.cashInBank))
		self.lines.append([['Cash Balance, '+month+' '+day+', '+year,'','','','',totalBalance],[0,0,0,0,0,1],[1,0,0,0,0,0]])

		self.commitLines(self.lines)

		self.popOAL()

		self.noteChunk.update()
		self.noteBlocks = self.noteChunk.noteBlocks
		print self.noteBlocks

	def getInflows(self,flowList):
		partialTotals={}
		totalInflows=0
		for i in flowList:
			name = i.getContents().nature.content
			amount = i.getContents().amount.content
			notes = i.note.content
			category = i.source.content.split(":")[0]
			if category=="COCPNote":
				name=i.getContents().event.content
			try:
				amount = float(amount)
			except:
				amount=0
			totalInflows+=amount
			if name in partialTotals:
				partialTotals[name][0]+=amount
			else:
				partialTotals[name]=[amount]
				partialTotals[name].append(notes)
		newLines = self.writeLinesPartial(partialTotals)
		return newLines,totalInflows

	def writeLinesPartial(self,partialTotalsList):
		total=0
		partialTotalsKeys = partialTotalsList.keys()
		partialTotalsKeys.sort()
		tempLines=[]
		for i in partialTotalsKeys:
			name = i
			amount = partialTotalsList[i][0]
			notes = partialTotalsList[i][1]
			total+=amount
			tempLines.append([[tab()+name,notes,'',floatToStr(amount),'',''],[],[]])

		if tempLines:
			if not self.firstInflow:
				tempLines[0][0][2]="P"
				self.firstInflow=1
			tempLines.append([[tab(2)+"Total",'','','','',floatToStr(total)],[0,0,0,0,0,0],[]])
			if self.lineTypes == "Outflow":
				tempLines[-1][0][-1]="("+tempLines[-1][0][-1]+")"
			if not self.firstInflowTotal:
				self.firstInflowTotal=1
				tempLines[-1][0][4]="P"
		return tempLines

	def commitLines(self,lines):
		self.cashFlowsText.clear()
		for i in lines:
			self.cashFlowsText.addRow(i[0],uls=i[1],bolds=i[2])

	def setUnderlines(self,lines):
		for i in range(0,len(lines)):
			if lines[i][0][0]==tab(2)+"Total":
				lines[i-1][1]=[0,0,0,1,0,0]
			if lines[i][0][0]==tab(3)+"Total Inflows" or lines[i][0][0]==tab(3)+"Total Outflows":
				lines[i-1][1]=[0,0,0,0,0,1]
		return lines

	def popOAL(self):
		oalInfo = getOAL(self.app)
		self.oalText['state']='normal'
		self.oalText.delete('1.0','end')
		if oalInfo!="":
			self.oalText.insert(self.oalText.index("end"),"\nOTHER ASSETS AND LIABILITIES\n")
			self.oalText.tag_add("bold",'1.0','end -1 line lineend')
			self.oalText.insert(self.oalText.index("end"),oalInfo)
			self.oalText.configure(height=int(self.oalText.index('end-1c').split('.')[0])+3)
		self.oalText['state']='disabled'


	def exportCallback(self):
		docBuilder = DocBuilder()

		# Write the 1st line of the header (council name)
		headerText = self.headerField.get(1.0, "1.end")
		docBuilder.createParagraph(headerText, "center", bold=True)

		# Write the 2nd line of the header (statement of cash flows)
		headerText = self.headerField.get(2.0, "2.end")
		docBuilder.createParagraph(headerText, "center", bold=True)

		# Write the 3rd line of the header (semester)
		headerText = self.headerField.get(3.0, "3.end")
		docBuilder.createParagraph(headerText, "center", bold=True)

		# Write the cash flows table
		self._export_CreateTable(self.lines, docBuilder)

		# Add some line breaks
		docBuilder.document.add_paragraph()

		# Write the OAL
		oalText = self.oalText.get(1.0, END)
		oalText = oalText.split("\n")
		docBuilder.document.add_paragraph().add_run(oalText[1]).bold = True
		for x in xrange(2, len(oalText)):
			docBuilder.document.add_paragraph(oalText[x])

		# Write the Notes header
		docBuilder.document.add_paragraph().add_run("NOTES TO FINANCIAL STATEMENT").bold = True

		# Write the Notes data
		for noteBlock in self.noteBlocks:
			if noteBlock.blockType == "text":
				# Write the heading
				newParagraph = docBuilder.document.add_paragraph()
				newParagraph.add_run(noteBlock.payload[0]).bold = True
				# Write the non-heading text
				textData = noteBlock.payload[1].split("\n")
				for line in textData:
					docBuilder.document.add_paragraph(line)
			elif noteBlock.blockType == "table":
				self._export_CreateTable(noteBlock.payload, docBuilder, 5)

		# Line break
		docBuilder.document.add_paragraph()

		# Write the statement of accountability
		docBuilder.document.add_paragraph().add_run("STATEMENT OF FINANCE COUNCILOR'S ACCOUNTABILITY").bold = True
		docBuilder.document.add_paragraph(""""I am hereunto affirming that to the best of my intentions, the Statement of Cash Flows and all its pertinent information fairly represent all the financial transactions that have transpired within the council, its constituents, and all relevant stakeholders for the semester ending (date), and in conformity with the standard accounting procedures as set forth and duly upheld by the University Student Council of UP Diliman." """)

		# Line break
		docBuilder.document.add_paragraph()
		docBuilder.document.add_paragraph()
		docBuilder.document.add_paragraph()

		# Write the signatures
		preparedBy = self.app.preparedBy.split("\n")
		notedBy = self.app.notedBy.split("\n")
		while len(preparedBy) < 3:
			preparedBy.append("")
		while len(notedBy) < 3:
			notedBy.append("")
		tableData = []
		tableData.append( [CellData("Prepared by:"), CellData("Noted by:")] )
		tableData.append( [CellData(""), CellData("")] )
		tableData.append( [CellData(""), CellData("")] )
		tableData.append( [CellData(preparedBy[0]), CellData(notedBy[0])] )
		tableData.append( [CellData(preparedBy[1]), CellData(notedBy[1])] )
		tableData.append( [CellData(preparedBy[2]), CellData(notedBy[2])] )
		docBuilder.createTable(2, tableData)

		# # Add signature images
		# run = docBuilder.table.cell(2, 0).add_paragraph().add_run()
		# run.add_picture('random.png', width=Inches(1))



		# Save the document
		fileName = 'StatementOfCashFlows_' + datetime.datetime.now().strftime("%I%M%p_%B%d_%Y") + '.docx'
		docBuilder.save(fileName)


	def _export_CreateTable(self, lineData, docBuilder, numberOfColumns=6):
		"""Reads self.lines and creates a docx table from it"""
		lines = list(lineData)

		# list of lists, where each entry is a list of CellData objects that represents a row
		tableData = [] 

		# iterate through the rows to construct tableData
		for line in lines:
			lineText = line[0]
			if not line[1]:
				line[1] = [0] * numberOfColumns
			if not line[2]:
				line[2] = [0] * numberOfColumns
			isUnderlined = line[1]
			isBold = line[2]

			rowData = []
			# iterate through the cells in a row to construct rowData
			for i in xrange(numberOfColumns):
				text = lineText[i]
				cellData = CellData(text)
				if (isUnderlined[i] == 1):
					cellData.tags.append("underline")
				if (isUnderlined[i] == 2):
					cellData.tags.append("double_underline")
				if (isBold[i] == 1):
					cellData.tags.append("bold")
				rowData.append(cellData)

			tableData.append(rowData)

		docBuilder.createTable(numberOfColumns, tableData)
		
		# manually set table column widths
		if (numberOfColumns == 6):
			docBuilder.columns[0].width = Inches(2.5)
			docBuilder.columns[1].width = Inches(1)
			docBuilder.columns[2].width = Inches(0.25)
			docBuilder.columns[3].width = Inches(1)
			docBuilder.columns[4].width = Inches(0.25)
			docBuilder.columns[5].width = Inches(1)
		elif (numberOfColumns == 5):
			docBuilder.columns[0].width = Inches(3.5)
			docBuilder.columns[1].width = Inches(0.25)
			docBuilder.columns[2].width = Inches(1)
			docBuilder.columns[3].width = Inches(0.25)
			docBuilder.columns[4].width = Inches(1)

if __name__=="__main__":
	root = Tk()
	app = StatementWindow(root,None,None)
	app.pack(fill=BOTH,expand=1)
	app.populateTree()
	app.mainloop()
